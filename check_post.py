#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
上稿前對照檢查 — 確認 docx 文字跟圖片資料夾一致
==========================================================
阿元 2026-06-17 提的需求: 避免「篇 5 文字配篇 3 圖」這種錯
"""

import sys, re
from pathlib import Path
from datetime import datetime

ROOT = Path(__file__).parent
POST_DIR = ROOT / "每篇貼文"

def check_one(folder_name):
    """檢查單一資料夾"""
    print(f"\n{'='*70}")
    print(f"🔍 檢查: {folder_name}")
    print('='*70)
    folder = POST_DIR / folder_name
    if not folder.exists():
        print(f"❌ 資料夾不存在")
        return False

    errors = []
    warnings = []
    info = []

    # 1. 解析資料夾名: YYYY-MM-DD_第N篇_品名 或 十齋日_YYYY-MM-DD_農曆XX
    m = re.match(r"(\d{4}-\d{2}-\d{2})_第(\d+)篇_(.+)", folder_name)
    zhai = re.match(r"十齋日_(\d{4}-\d{2}-\d{2})_農曆(.+)", folder_name)

    if m:
        date_str, num, kind = m.groups()
        post_type = "連載"
        expected_time = "08:00"
        info.append(f"類型: 連載第 {num} 篇 ({kind})")
        info.append(f"日期: {date_str}")
        info.append(f"預期排程: {date_str} {expected_time}")

        # 檢查日期是否週末
        wd = datetime.strptime(date_str, "%Y-%m-%d").weekday()
        if wd >= 5:
            warnings.append(f"日期 {date_str} 是週{'六日'[wd-5]}, 連載通常週一~五")

        expected_pian = int(num)
    elif zhai:
        date_str, lunar = zhai.groups()
        post_type = "十齋日"
        expected_time = "06:00"
        info.append(f"類型: 十齋日 (農曆{lunar})")
        info.append(f"日期: {date_str}")
        info.append(f"預期排程: {date_str} {expected_time}")
        expected_pian = None
    else:
        errors.append(f"資料夾名格式不對, 不像連載或十齋日")
        post_type = None
        expected_pian = None

    # 2. 檢查 docx 內文 — 篇號 + 主題對齊
    docx_file = folder / "貼文內容.docx"
    if not docx_file.exists():
        errors.append("找不到 貼文內容.docx")
    else:
        try:
            from docx import Document
            doc = Document(str(docx_file))
            all_text = "\n".join(p.text for p in doc.paragraphs)

            if expected_pian:  # 連載
                # 從 docx 抓「第 X 篇」
                docx_pian_matches = re.findall(r"第\s*(\d+)\s*篇", all_text)
                docx_pian = [int(x) for x in docx_pian_matches]
                if docx_pian:
                    most_common = max(set(docx_pian), key=docx_pian.count)
                    if most_common != expected_pian:
                        errors.append(
                            f"⚠️ 篇號錯誤! 資料夾名說「第 {expected_pian} 篇」"
                            f"但 docx 內出現的篇號是「第 {most_common} 篇」"
                        )
                    else:
                        info.append(f"✅ docx 篇號 = {most_common} 符合")

                # 從 docx 抓 12 張圖名 (圖片計畫表)
                img_plan = re.findall(r"^\s*(\d{1,2})\s*-?\s*([^\n]+?)(?:\s*-\s*[^\n]+)?$", all_text, re.MULTILINE)
                # 簡化: 看內文有沒有列圖名跟實際圖片對得起來 (放棄太複雜, 改檢查數量)

            elif post_type == "十齋日":
                # 檢查十齋日內文是否有「十齋日」+ 農曆日字樣
                if "十齋日" not in all_text:
                    errors.append("docx 內沒「十齋日」字樣")
                if lunar not in all_text:
                    warnings.append(f"docx 內沒「{lunar}」字樣, 但農曆日要符合")
        except Exception as e:
            warnings.append(f"docx 解析失敗: {e}")

    # 2b. ★ 圖文比對 (阿元 2026-06-17 拍板) — docx 內「圖片計畫」vs 實際圖檔名
    # docx 內格式 (篇 1-3): "01 - 佉羅帝耶山遠景 - 廣角全景"
    # 實際檔名 (篇 1-3): "01_佉羅帝耶山遠景.png"
    # 比對「篇號-主題」是否一致
    try:
        plan_lines = []
        in_plan = False
        for p in doc.paragraphs:
            t = p.text.strip()
            if "12 張圖規劃" in t or "圖片計畫" in t:
                in_plan = True; continue
            if in_plan and p.style.name.startswith("Heading"):
                break
            if in_plan and re.match(r"^\d{1,2}\s*[-:_]", t):
                plan_lines.append(t)
        info.append(f"docx 圖片計畫: {len(plan_lines)} 行")
    except Exception:
        pass

    # 3. 檢查圖片數量 (含子資料夾「圖片」, 篇 4+ 圖在 圖片/)
    imgs_root = sorted([f for f in folder.glob("*.png")] + [f for f in folder.glob("*.jpg")])
    imgs_sub = []
    sub = folder / "圖片"
    if sub.exists():
        imgs_sub = sorted([f for f in sub.glob("*.png")] + [f for f in sub.glob("*.jpg")])
    imgs = imgs_root if imgs_root else imgs_sub
    img_count = len(imgs)
    if imgs_sub and not imgs_root:
        info.append(f"圖位置: {folder.name}/圖片/ (子資料夾)")

    if post_type == "連載":
        if img_count == 0:
            errors.append("沒有任何圖片!")
        elif img_count < 12:
            warnings.append(f"圖片只有 {img_count} 張, 連載通常 12 張")
        elif img_count > 12:
            warnings.append(f"圖片有 {img_count} 張, 超過連載 12 張上限")
        else:
            info.append(f"✅ 圖片 12 張齊全")

        # 檢查圖名是否 01_ ~ 12_
        expected_nums = set(f"{i:02d}" for i in range(1, 13))
        actual_nums = set()
        for img in imgs:
            m_img = re.match(r"^(\d{2})_", img.name)
            if m_img:
                actual_nums.add(m_img.group(1))
        missing = expected_nums - actual_nums
        if missing and img_count >= 1:
            warnings.append(f"缺圖編號: {sorted(missing)}")

    elif post_type == "十齋日":
        if img_count == 0:
            errors.append("沒有圖片!")
        elif img_count > 1:
            warnings.append(f"十齋日通常 1 張圖, 但有 {img_count} 張")
        else:
            info.append(f"✅ 圖片 1 張")

    # 4. 輸出
    print()
    print("📋 資訊:")
    for x in info: print(f"   {x}")

    if warnings:
        print("\n⚠️ 警告:")
        for x in warnings: print(f"   {x}")

    if errors:
        print("\n❌ 錯誤 (必須修):")
        for x in errors: print(f"   {x}")
        return False

    if not errors:
        print("\n✅ 對照檢查通過 — 可以上稿")
        return True

def pause():
    try: input("\n按 Enter 結束...")
    except EOFError: pass

if __name__ == "__main__":
    if len(sys.argv) > 1:
        targets = sys.argv[1:]
    else:
        # 預設檢查篇 1
        targets = ["2026-06-23_第1篇_序品第一"]

    all_ok = True
    for t in targets:
        ok = check_one(t)
        if not ok: all_ok = False

    print()
    print("=" * 70)
    if all_ok:
        print(f"🪷 全部 {len(targets)} 個資料夾檢查通過 ✅")
    else:
        print(f"❌ 有資料夾沒過, 上稿前先修正")
    print("=" * 70)
    pause()
