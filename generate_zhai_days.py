"""
十齋日 docx 生成腳本
阿地 4 號 2026-05-20

★ 規則 (阿元 2026-05-15 project_instructions):
- 農曆 1, 8, 14, 15, 18, 23, 24, 28, 29, 30 為十齋日
- 每日生成 1 個 docx, 放到 每篇貼文/十齋日_YYYY-MM-DD_農曆XXX/
- 早上 6:00 上刊
- 不跳過週末 / 國定假日 (跟 100 篇連載分開)
- 固定文字 (《地藏菩薩本願經》原文):
  「復次普廣。若未來世眾生。於月一日。八日。...能於是十齋日。對佛菩薩。
   諸賢聖像前。讀是經一徧。東西南北。百由旬內。無諸災難。」
"""
from datetime import date, timedelta
from pathlib import Path
from lunardate import LunarDate
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = PROJECT_ROOT / "每篇貼文"

# 農曆十齋日
ZHAI_DAYS = [1, 8, 14, 15, 18, 23, 24, 28, 29, 30]

# 起訖日 (從今天到 2027 年底)
START = date(2026, 5, 20)
END = date(2027, 12, 31)

# 固定文字 (《地藏菩薩本願經》原文 — 從 project_instructions)
FIXED_SUTRA_TEXT = """復次普廣。若未來世眾生。於月一日。八日。十四日。十五日。十八日。二十三。二十四。二十八。二十九日。乃至三十日。是諸日等。諸罪結集。定其輕重。

南閻浮提眾生。舉止動念。無不是業。無不是罪。何況恣情殺害竊盜邪淫妄語。百千罪狀。

能於是十齋日。對佛菩薩。諸賢聖像前。讀是經一徧（ㄅㄧㄢˋ）。東西南北。百由旬內。無諸災難。當此居家。若長若幼。現在未來。百千歲中。永離惡趣。

能於十齋日。每轉一徧（ㄅㄧㄢˋ）。現世令此居家。無諸橫病。衣食豐溢。"""


def lunar_day_name(lunar_day):
    """農曆日轉中文 (初一, 初二, 初十, 十一, ...)"""
    chinese_nums = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if lunar_day <= 10:
        return f"初{chinese_nums[lunar_day]}"
    elif lunar_day < 20:
        return f"十{chinese_nums[lunar_day - 10]}" if lunar_day > 10 else "十"
    elif lunar_day == 20:
        return "二十"
    elif lunar_day < 30:
        return f"二十{chinese_nums[lunar_day - 20]}"
    else:
        return "三十"


def weekday_chinese(d):
    names = ["一", "二", "三", "四", "五", "六", "日"]
    return names[d.weekday()]


def find_zhai_days(start, end):
    """從 start 到 end 找所有十齋日, 回傳 [(solar_date, lunar_day, lunar_str), ...]"""
    days = []
    cur = start
    while cur <= end:
        ld = LunarDate.fromSolarDate(cur.year, cur.month, cur.day)
        if ld.day in ZHAI_DAYS:
            days.append((cur, ld.day, lunar_day_name(ld.day)))
        cur += timedelta(days=1)
    return days


def generate_zhai_docx(solar_date, lunar_day, lunar_str):
    """生成單篇十齋日 docx"""
    date_str = solar_date.strftime("%Y-%m-%d")
    weekday = weekday_chinese(solar_date)
    folder_name = f"十齋日_{date_str}_農曆{lunar_str}"
    folder = OUTPUT_DIR / folder_name
    folder.mkdir(parents=True, exist_ok=True)

    docx_path = folder / "貼文內容.docx"

    doc = Document()

    # Heading 1
    doc.add_heading(f"🌸 今日 — 十齋日 · 農曆{lunar_str}", level=1)

    # Heading 2 + FB 貼文
    doc.add_heading("📱 FB 貼文文字", level=2)

    # FB 貼文段落
    paragraphs = [
        f"🪷【十齋日 · 農曆{lunar_str}】",
        "",
        "親愛的朋友, 早安 🙏",
        "",
        f"今天是農曆{lunar_str}日 — 「十齋日」.",
        "",
        "★ 什麼是「十齋日」?",
        "",
        "《地藏菩薩本願經》中, 佛陀親口告訴普廣菩薩 — ",
        "農曆每月的「一、八、十四、十五、十八、二十三、二十四、二十八、二十九、三十」",
        "這 10 天, 諸罪會被「結集定其輕重」(諸佛菩薩會審視眾生的善惡).",
        "",
        "所以這 10 天是修行最重要的日子!",
        "",
        "───────────",
        "",
        "📖 經文 (《地藏菩薩本願經》原文):",
        "",
        FIXED_SUTRA_TEXT,
        "",
        "───────────",
        "",
        "✨ 經文解釋 (淺顯版):",
        "",
        "🌸 在十齋日這 10 天 — ",
        "🌸 對佛菩薩 — 讀《地藏菩薩本願經》一遍",
        "🌸 你家方圓百由旬 (大範圍) — 無災難",
        "🌸 你跟家人 — 100 年到 1000 年內, 永遠不墮惡道",
        "🌸 你的家 — 沒有意外的疾病, 衣食豐足!",
        "",
        "───────────",
        "",
        "💡 今天怎麼修?",
        "",
        "1. 🪷 早起 — 心想著《地藏菩薩本願經》",
        "2. 🪷 讀經一遍 — 不必很久, 心誠就好",
        "3. 🪷 不殺生 — 一天試著吃素",
        "4. 🪷 不惡口 — 一天說好話",
        "5. 🪷 行善 — 做一件幫助他人的小事",
        "",
        "───────────",
        "",
        "📚 免費索取《地藏菩薩本願經》紙本經書",
        "",
        "私訊我們「我要索取地藏經」, 將免費寄送精美紙本經書給您(含郵資)🌸",
        "",
        f"#地藏菩薩 #十齋日 #農曆{lunar_str} #地藏菩薩本願經",
    ]
    for p in paragraphs:
        doc.add_paragraph(p)

    # Heading 2 + 上刊資訊
    doc.add_heading("📋 上刊資訊", level=2)
    doc.add_paragraph(
        f"上刊日期: {date_str} (週{weekday}) · 早 06:00 · 地藏菩薩本行經"
    )
    doc.add_paragraph("★ 十齋日 — 比 100 篇連載的 07:00 早 1 小時上刊")

    doc.save(str(docx_path))
    return docx_path


def main():
    print(f"找十齋日 {START} ~ {END}")
    days = find_zhai_days(START, END)
    print(f"共 {len(days)} 個十齋日\n")

    ok = 0
    for solar_date, lunar_day, lunar_str in days:
        try:
            p = generate_zhai_docx(solar_date, lunar_day, lunar_str)
            print(f"OK {solar_date} 農曆{lunar_str}")
            ok += 1
        except Exception as e:
            print(f"FAIL {solar_date}: {e}")
    print(f"\n總計: OK={ok}, 全 {len(days)} 個十齋日")


if __name__ == "__main__":
    main()
